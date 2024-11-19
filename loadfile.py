import base64
import os
import re
import tempfile
from gridfs import GridFS
import pypandoc
import pytesseract
from pdf2image import convert_from_path
import streamlit as st
from docx import Document
import spacy
from pymongo import MongoClient
from PyPDF2 import PdfReader
from streamlit_js_eval import streamlit_js_eval

# Load spaCy's English model for NLP tasks
nlp = spacy.load("en_core_web_sm")

# Retrieve MongoDB URI from environment variables
MONGODB_URI = os.getenv("MONGODB_URI")

# Configure MongoDB connection
client = MongoClient(MONGODB_URI)
db = client['resume_match_ai']
fs = GridFS(db)

# Refresh the page function
def refresh_page():
    """ Refreshes the page by reloading the session state. """
    st.session_state.show_job_details = False
    st.session_state.selected_job_id = None
    streamlit_js_eval(js_expressions="parent.window.location.reload()")

# OCR function with retry mechanism for scanned PDFs
def ocr_pdf(file, retries=3):
    """Perform OCR on a PDF file to extract text with retries."""
    try:
        images = convert_from_path(file, 300)
        text = ""
        for image in images:
            text += pytesseract.image_to_string(image)
        return text
    except Exception as e:
        if retries > 0:
            return ocr_pdf(file, retries - 1)
        return f"Error performing OCR: {str(e)}"

# Extract text from DOCX file
def extract_text_from_docx(filepath):
    """Extract text from a DOCX file, including tables."""
    try:
        doc = Document(filepath)
        text = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

# Extract text from PDF with enhanced logging
def extract_text_from_pdf(filepath):
    """Extract text from a PDF file with PyPDF2, fallback to OCR if empty pages."""
    text = []
    try:
        pdfReader = PdfReader(filepath)
        for i, page in enumerate(pdfReader.pages):
            page_text = page.extract_text()
            if page_text:
                text.append(page_text + "\n\n")
            else:
                text.append(f"[OCR Page {i+1}]\n" + ocr_pdf(filepath))  # Fallback to OCR
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"
    return "\n".join(text)

# Extract text from TXT file
def extract_text_from_txt(filepath):
    """Extract text from a TXT file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        return f"Error extracting text from TXT: {str(e)}"

# Extract text from file, with exception handling for unsupported files
def extract_text_from_file(uploaded_file):
    """Extract text from an uploaded file (DOCX, PDF, or TXT)."""
    try:
        if uploaded_file.name.endswith('.docx'):
            return extract_text_from_docx(uploaded_file)
        elif uploaded_file.name.endswith('.pdf'):
            return extract_text_from_pdf(uploaded_file)
        elif uploaded_file.name.endswith('.txt'):
            return extract_text_from_txt(uploaded_file)
        else:
            return "Unsupported file format."
    except Exception as e:
        return f"Error extracting text: {str(e)}"

# Clean text to remove extra formatting
def clean_text(text):
    """Clean extracted text to remove unnecessary formatting and common artifacts."""
    text = re.sub(r'\n+', ' ', text)  # Replace multiple newlines
    text = re.sub(r'\s{2,}', ' ', text)  # Replace multiple spaces
    text = re.sub(r'https?://\S+', '', text)  # Remove URLs
    return text.strip()

# Extract entities with improved skills detection
def extract_entities(text):
    """Extract named entities using spaCy and regex for custom entities."""
    doc = nlp(text)
    entities = {
        "names": [],
        "skills": [],
        "organizations": [],
        "dates": [],
    }

    for ent in doc.ents:
        if ent.label_ == "PERSON":
            entities["names"].append(ent.text)
        elif ent.label_ == "ORG":
            entities["organizations"].append(ent.text)
        elif ent.label_ == "DATE":
            entities["dates"].append(ent.text)

    # Regex-based skill extraction if 'SKILL' is not in spaCy model
    skill_pattern = re.compile(r'\b(Java|Python|SQL|C\+\+|Machine Learning|Data Science|NLP|AWS)\b', re.IGNORECASE)
    entities["skills"].extend(skill_pattern.findall(text))
    
    return entities

# Process files with improved error handling and logging
def process_files(uploaded_files):
    """Process uploaded files, extracting text and entities with error handling."""
    results = []
    for uploaded_file in uploaded_files:
        extracted_text = extract_text_from_file(uploaded_file)
        if extracted_text and "Error" not in extracted_text:
            cleaned_text = clean_text(extracted_text)
            entities = extract_entities(cleaned_text)
            results.append({
                "file_name": uploaded_file.name,
                "cleaned_text": cleaned_text,
                "entities": entities
            })
        else:
            results.append({
                "file_name": uploaded_file.name,
                "error": extracted_text  # Shows specific error for easier debugging
            })
    return results
    
def extract_job_description(uploaded_job):
    """Extract job description from the uploaded file."""
    if uploaded_job.type == 'application/pdf':
        return extract_text_from_file(uploaded_job)
    elif uploaded_job.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        return extract_text_from_file(uploaded_job)
    else:
        st.error("Unsupported file format")
        return None

def handle_input(input_type):
    """Handle file or textarea input based on the input type."""
    uploaded_job = None
    job_description_text = None

    if input_type == 'file':
        uploaded_job = st.file_uploader('Upload Job Description:', type=['pdf', 'docx'], 
                                        key="job_file_uploader")
    elif input_type == 'textarea':
        job_description_text = st.text_area("Or write the job description here:")

    return uploaded_job, job_description_text


def set_input_type(input_type):
    st.session_state.use_file = input_type == 'file'
    st.session_state.use_textarea = input_type == 'textarea'

# Function to refresh page by clearing input usage states
def cancel_input():
    st.session_state.use_file = False
    st.session_state.use_textarea = False
    refresh_page()

def display_field(label, data):
    st.markdown(f"**{label}:**")
    if data and data[0] != "N/A":
        for item in data:
            st.markdown(f"- {item.strip()}")
    else:
        st.write("N/A")

def format_name(name):
    """Format candidate name for display."""
    return ' '.join([word.capitalize() for word in name.split()])
